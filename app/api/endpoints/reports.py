"""Экспорт отчётов в Excel и Word."""

from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from typing import Annotated, Any
from urllib.parse import quote
from uuid import UUID

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from supabase import Client

from app.api.deps import require_master_or_admin, require_report_access
from app.core.db_utils import execute_supabase, first_row
from app.core.helpers import normalize_join
from app.core.supabase import get_supabase_client
from app.models.schemas import CurrentUser

router = APIRouter(prefix="/reports", tags=["reports"])

THEME_HEADER_BG = "#212529"
THEME_HEADER_FG = "#FFFFFF"
BRAND_TITLE = "ОАО «БААЗ»"
DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _excel_response(buffer: BytesIO, filename: str) -> StreamingResponse:
    buffer.seek(0)
    safe_filename = filename or "report.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_filename)}"},
    )


def _cell(value: Any) -> str:
    """Приводит значение ячейки к строке для pandas/openpyxl."""
    if value is None:
        return ""
    if isinstance(value, UUID):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, (int, float, bool)):
        return str(value)
    if isinstance(value, dict):
        return ""
    if isinstance(value, list):
        return ", ".join(_cell(item) for item in value if item is not None)
    return str(value)


def _normalize_join(value: Any) -> dict[str, Any]:
    """Приводит вложенный join Supabase к одному словарю."""
    return normalize_join(value)


def _rows_to_dataframe(rows: list[dict[str, Any]], columns: list[str]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=columns)
    flat_rows = [{col: _cell(row.get(col, "")) for col in columns} for row in rows]
    return pd.DataFrame(flat_rows, columns=columns)


def _column_widths(df: pd.DataFrame) -> list[int]:
    widths: list[int] = []
    for col in df.columns:
        values = [str(col)] + [str(v) if v is not None else "" for v in df[col].tolist()]
        widths.append(min(max(len(v) for v in values) + 2, 60))
    return widths


def _write_styled_sheet(
    writer: pd.ExcelWriter,
    sheet_name: str,
    df: pd.DataFrame,
    *,
    report_title: str,
    subtitle: str | None = None,
) -> None:
    """Записывает лист с заголовком, шапкой таблицы и границами ячеек."""
    workbook = writer.book
    worksheet = workbook.add_worksheet(sheet_name[:31])
    writer.sheets[sheet_name[:31]] = worksheet

    ncols = max(len(df.columns), 1) - 1

    title_fmt = workbook.add_format({"bold": True, "font_size": 14, "valign": "vcenter"})
    date_fmt = workbook.add_format({"italic": True, "font_color": "#495057"})
    header_fmt = workbook.add_format(
        {
            "bold": True,
            "bg_color": THEME_HEADER_BG,
            "font_color": THEME_HEADER_FG,
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
        }
    )
    cell_fmt = workbook.add_format({"border": 1, "valign": "vcenter", "text_wrap": True})

    row = 0
    worksheet.merge_range(row, 0, row, ncols, report_title, title_fmt)
    row += 1
    worksheet.merge_range(row, 0, row, ncols, f"Дата формирования: {date.today().isoformat()}", date_fmt)
    row += 1
    if subtitle:
        worksheet.merge_range(row, 0, row, ncols, subtitle, date_fmt)
        row += 1
    row += 1  # пустая строка перед таблицей

    header_row = row
    for col_idx, col_name in enumerate(df.columns):
        worksheet.write(header_row, col_idx, col_name, header_fmt)

    for data_row_idx, row_values in enumerate(df.itertuples(index=False), start=header_row + 1):
        for col_idx, value in enumerate(row_values):
            worksheet.write(data_row_idx, col_idx, value, cell_fmt)

    for col_idx, width in enumerate(_column_widths(df)):
        worksheet.set_column(col_idx, col_idx, width)


def _build_excel_workbook(
    sheets: list[tuple[str, pd.DataFrame, str, str | None]],
) -> BytesIO:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        for sheet_name, df, title, subtitle in sheets:
            _write_styled_sheet(writer, sheet_name, df, report_title=title, subtitle=subtitle)
    return buffer


def _write_inventory_combined_sheet(
    writer: pd.ExcelWriter,
    *,
    meta_rows: list[tuple[str, str]],
    tools_df: pd.DataFrame,
    warehouse_name: str,
    location_name: str,
) -> None:
    """Один лист: сводка сверху, таблица инструментов с 10-й строки Excel."""
    workbook = writer.book
    sheet_name = "Ведомость"
    worksheet = workbook.add_worksheet(sheet_name)
    writer.sheets[sheet_name] = worksheet

    ncols = max(len(INVENTORY_COLUMNS), 2) - 1
    title_fmt = workbook.add_format({"bold": True, "font_size": 14, "valign": "vcenter"})
    date_fmt = workbook.add_format({"italic": True, "font_color": "#495057"})
    meta_label_fmt = workbook.add_format({"bold": True, "border": 1, "valign": "vcenter"})
    meta_value_fmt = workbook.add_format({"border": 1, "valign": "vcenter"})
    header_fmt = workbook.add_format(
        {
            "bold": True,
            "bg_color": THEME_HEADER_BG,
            "font_color": THEME_HEADER_FG,
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
        }
    )
    cell_fmt = workbook.add_format({"border": 1, "valign": "vcenter", "text_wrap": True})

    row = 0
    worksheet.merge_range(
        row,
        0,
        row,
        ncols,
        f"{BRAND_TITLE} — Ведомость остатков инструмента",
        title_fmt,
    )
    row += 1
    worksheet.merge_range(row, 0, row, ncols, f"Дата формирования: {date.today().isoformat()}", date_fmt)
    row += 1
    worksheet.merge_range(
        row,
        0,
        row,
        ncols,
        f"Склад: {warehouse_name} · Подразделение: {location_name or '—'}",
        date_fmt,
    )
    row += 2

    for label, value in meta_rows:
        worksheet.write(row, 0, label, meta_label_fmt)
        worksheet.write(row, 1, value, meta_value_fmt)
        row += 1

    header_row = INVENTORY_TABLE_START_ROW
    for col_idx, col_name in enumerate(tools_df.columns):
        worksheet.write(header_row, col_idx, col_name, header_fmt)

    for data_row_idx, row_values in enumerate(tools_df.itertuples(index=False), start=header_row + 1):
        for col_idx, value in enumerate(row_values):
            worksheet.write(data_row_idx, col_idx, value, cell_fmt)

    meta_df = pd.DataFrame(meta_rows, columns=["Параметр", "Значение"])
    col0_width = max(_column_widths(meta_df)[0], _column_widths(tools_df)[0] if len(tools_df.columns) else 12)
    worksheet.set_column(0, 0, col0_width)
    for col_idx, width in enumerate(_column_widths(tools_df)):
        if col_idx == 0:
            continue
        worksheet.set_column(col_idx, col_idx, width)


def _build_inventory_workbook(
    meta_rows: list[tuple[str, str]],
    tools_df: pd.DataFrame,
    warehouse_name: str,
    location_name: str,
) -> BytesIO:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        _write_inventory_combined_sheet(
            writer,
            meta_rows=meta_rows,
            tools_df=tools_df,
            warehouse_name=warehouse_name,
            location_name=location_name,
        )
    return buffer


def _extract_writeoff_comment(tool: dict[str, Any]) -> str:
    lines_raw = tool.get("tms_requisition_lines") or []
    if isinstance(lines_raw, dict):
        lines_raw = [lines_raw]
    for line in reversed(lines_raw):
        if line.get("status") == "returned" and line.get("condition_on_return"):
            return _cell(line.get("condition_on_return"))
    return ""


def _tool_writeoff_date(tool: dict[str, Any]) -> date | None:
    raw = tool.get("last_check")
    if not raw:
        return None
    try:
        return date.fromisoformat(str(raw)[:10])
    except ValueError:
        return None


def _flatten_writeoff_tool(tool: dict[str, Any]) -> dict[str, str]:
    tool_type = _normalize_join(tool.get("tms_tool_types"))
    last_check = tool.get("last_check")

    return {
        "Инв. номер": _cell(tool.get("inventory_number")),
        "Серийный номер": _cell(tool.get("serial_number")),
        "Модель": _cell(tool_type.get("model_name")),
        "Дата списания": _cell(last_check)[:10] if last_check else "",
        "Причина (комментарий)": _extract_writeoff_comment(tool),
    }


def _filter_scrapped_tools_by_period(
    tools: list[dict[str, Any]],
    date_from: date,
    date_to: date,
) -> list[dict[str, Any]]:
    """Инструменты со status=scrapped, дата списания (last_check) в заданном периоде."""
    filtered: list[dict[str, Any]] = []
    for tool in tools:
        if tool.get("status") != "scrapped":
            continue
        writeoff_date = _tool_writeoff_date(tool)
        if writeoff_date is not None and date_from <= writeoff_date <= date_to:
            filtered.append(tool)
    return filtered


WRITEOFF_TABLE_START_ROW = 9
WRITEOFF_EMPTY_MESSAGE = "За указанный период списаний не зафиксировано"


def _write_writeoffs_combined_sheet(
    writer: pd.ExcelWriter,
    *,
    meta_rows: list[tuple[str, str]],
    tools_df: pd.DataFrame,
    period_label: str,
    empty: bool,
) -> None:
    """Один лист: сводка + таблица списаний с 10-й строки или сообщение об отсутствии данных."""
    workbook = writer.book
    sheet_name = "Списания"
    worksheet = workbook.add_worksheet(sheet_name)
    writer.sheets[sheet_name] = worksheet

    ncols = max(len(WRITEOFF_COLUMNS), 2) - 1
    title_fmt = workbook.add_format({"bold": True, "font_size": 14, "valign": "vcenter"})
    date_fmt = workbook.add_format({"italic": True, "font_color": "#495057"})
    meta_label_fmt = workbook.add_format({"bold": True, "border": 1, "valign": "vcenter"})
    meta_value_fmt = workbook.add_format({"border": 1, "valign": "vcenter"})
    header_fmt = workbook.add_format(
        {
            "bold": True,
            "bg_color": THEME_HEADER_BG,
            "font_color": THEME_HEADER_FG,
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
        }
    )
    cell_fmt = workbook.add_format({"border": 1, "valign": "vcenter", "text_wrap": True})
    empty_fmt = workbook.add_format({"italic": True, "font_color": "#6c757d", "font_size": 12})

    row = 0
    worksheet.merge_range(row, 0, row, ncols, f"{BRAND_TITLE} — Отчёт по списаниям", title_fmt)
    row += 1
    worksheet.merge_range(row, 0, row, ncols, f"Дата формирования: {date.today().isoformat()}", date_fmt)
    row += 1
    worksheet.merge_range(row, 0, row, ncols, period_label, date_fmt)
    row += 2

    for label, value in meta_rows:
        worksheet.write(row, 0, label, meta_label_fmt)
        worksheet.write(row, 1, value, meta_value_fmt)
        row += 1

    if empty:
        worksheet.merge_range(WRITEOFF_TABLE_START_ROW, 0, WRITEOFF_TABLE_START_ROW, ncols, WRITEOFF_EMPTY_MESSAGE, empty_fmt)
        worksheet.set_column(0, 0, len(WRITEOFF_EMPTY_MESSAGE) + 2)
        return

    header_row = WRITEOFF_TABLE_START_ROW
    for col_idx, col_name in enumerate(tools_df.columns):
        worksheet.write(header_row, col_idx, col_name, header_fmt)

    for data_row_idx, row_values in enumerate(tools_df.itertuples(index=False), start=header_row + 1):
        for col_idx, value in enumerate(row_values):
            worksheet.write(data_row_idx, col_idx, value, cell_fmt)

    for col_idx, width in enumerate(_column_widths(tools_df)):
        worksheet.set_column(col_idx, col_idx, width)


def _build_writeoffs_workbook(
    meta_rows: list[tuple[str, str]],
    tools_df: pd.DataFrame,
    period_label: str,
    *,
    empty: bool,
) -> BytesIO:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        _write_writeoffs_combined_sheet(
            writer,
            meta_rows=meta_rows,
            tools_df=tools_df,
            period_label=period_label,
            empty=empty,
        )
    return buffer


def _fetch_warehouse(supabase: Client, warehouse_id: UUID) -> dict[str, Any]:
    response = execute_supabase(
        lambda: supabase.table("tms_warehouses")
        .select("id, name, tms_locations(name)")
        .eq("id", str(warehouse_id))
        .execute()
    )
    return first_row(response, detail="Склад не найден")


def _flatten_inventory_tool(tool: dict[str, Any]) -> dict[str, str]:
    tool_type = _normalize_join(tool.get("tms_tool_types"))
    category = _normalize_join(tool_type.get("tms_tool_categories"))
    last_check_raw = tool.get("last_check")
    last_check = _cell(last_check_raw)[:10] if last_check_raw else ""

    status_labels = {
        "available": "Доступен",
        "in_use": "В работе",
        "maintenance": "На обслуживании",
        "scrapped": "Списан",
    }
    status = tool.get("status") or ""

    return {
        "Инв. номер": _cell(tool.get("inventory_number")),
        "Серийный номер": _cell(tool.get("serial_number")),
        "Модель": _cell(tool_type.get("model_name")),
        "Категория": _cell(category.get("name")),
        "Статус": status_labels.get(status, _cell(status)),
        "Износ": _cell(tool.get("wear_count")),
        "Дата последней поверки": last_check,
    }


INVENTORY_TABLE_START_ROW = 9  # 0-based: Excel row 10
INVENTORY_COLUMNS = [
    "Инв. номер",
    "Серийный номер",
    "Модель",
    "Категория",
    "Статус",
    "Износ",
    "Дата последней поверки",
]

WRITEOFF_COLUMNS = [
    "Инв. номер",
    "Серийный номер",
    "Модель",
    "Дата списания",
    "Причина (комментарий)",
]


@router.get("/export/inventory/{warehouse_id}")
def export_inventory(
    warehouse_id: UUID,
    supabase: Annotated[Client, Depends(get_supabase_client)],
    current_user: Annotated[CurrentUser, Depends(require_report_access)],
) -> StreamingResponse:
    """Ведомость остатков инструментов на складе."""
    warehouse = _fetch_warehouse(supabase, warehouse_id)
    warehouse_id_str = str(warehouse_id)

    if current_user.role == "clerk" and current_user.warehouse_id:
        if str(current_user.warehouse_id) != str(warehouse_id):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Доступ к другому складу запрещён")

    tools_response = execute_supabase(
        lambda: supabase.table("tms_tools")
        .select(
            "inventory_number, serial_number, status, wear_count, last_check, "
            "tms_tool_types(model_name, tms_tool_categories(name))"
        )
        .eq("warehouse_id", warehouse_id_str)
        .order("inventory_number")
        .execute()
    )

    rows = [_flatten_inventory_tool(tool) for tool in tools_response.data or []]
    location = _normalize_join(warehouse.get("tms_locations"))
    warehouse_name = _cell(warehouse.get("name"))

    df = _rows_to_dataframe(rows, INVENTORY_COLUMNS)
    location_name = _cell(location.get("name"))
    meta_rows = [
        ("Склад", warehouse_name),
        ("Подразделение", location_name or "—"),
        ("Всего позиций", _cell(len(rows))),
    ]

    buffer = _build_inventory_workbook(
        meta_rows=meta_rows,
        tools_df=df,
        warehouse_name=warehouse_name,
        location_name=location_name,
    )

    safe_name = _cell(warehouse.get("name")).replace(" ", "_")[:40] or "warehouse"
    return _excel_response(buffer, f"inventory_{safe_name}.xlsx")


@router.get("/export/write-offs")
def export_write_offs(
    supabase: Annotated[Client, Depends(get_supabase_client)],
    _: Annotated[CurrentUser, Depends(require_master_or_admin)],
    date_from: date = Query(..., description="Начало периода"),
    date_to: date = Query(..., description="Конец периода"),
) -> StreamingResponse:
    """Отчёт по списанному инструменту за период."""
    if date_from > date_to:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="date_from не может быть позже date_to")

    response = execute_supabase(
        lambda: supabase.table("tms_tools")
        .select(
            "inventory_number, serial_number, status, last_check, "
            "tms_tool_types(model_name), "
            "tms_requisition_lines(condition_on_return, status)"
        )
        .eq("status", "scrapped")
        .order("last_check")
        .execute()
    )

    scrapped_tools = _filter_scrapped_tools_by_period(response.data or [], date_from, date_to)
    rows = [_flatten_writeoff_tool(tool) for tool in scrapped_tools]
    df = _rows_to_dataframe(rows, WRITEOFF_COLUMNS)
    period_label = f"Период: {date_from.isoformat()} — {date_to.isoformat()}"
    meta_rows = [
        ("Период с", date_from.isoformat()),
        ("Период по", date_to.isoformat()),
        ("Записей", _cell(len(rows))),
    ]

    buffer = _build_writeoffs_workbook(
        meta_rows=meta_rows,
        tools_df=df,
        period_label=period_label,
        empty=not rows,
    )

    return _excel_response(buffer, f"write_offs_{date_from}_{date_to}.xlsx")


def _fetch_scrapped_tool(supabase: Client, tool_id: UUID) -> dict[str, Any]:
    """Загружает списанный инструмент с данными для акта."""
    response = execute_supabase(
        lambda: supabase.table("tms_tools")
        .select(
            "id, inventory_number, serial_number, status, last_check, warehouse_id, "
            "tms_tool_types(model_name), "
            "tms_warehouses(name, tms_locations(name)), "
            "tms_requisition_lines(condition_on_return, status)"
        )
        .eq("id", str(tool_id))
        .limit(1)
        .execute()
    )
    tool = first_row(response, detail="Инструмент не найден")
    if tool.get("status") != "scrapped":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Акт списания доступен только для инструментов со статусом «Списан»",
        )
    return tool


def _build_writeoff_docx(tool: dict[str, Any]) -> BytesIO:
    tool_type = _normalize_join(tool.get("tms_tool_types"))
    warehouse = _normalize_join(tool.get("tms_warehouses"))
    location = _normalize_join(warehouse.get("tms_locations"))
    reason = _extract_writeoff_comment(tool) or "—"
    writeoff_date = _cell(tool.get("last_check"))[:10] or date.today().isoformat()
    model_name = _cell(tool_type.get("model_name")) or "—"
    inventory_number = _cell(tool.get("inventory_number")) or "—"
    serial_number = _cell(tool.get("serial_number")) or "—"
    place = _cell(location.get("name")) or _cell(warehouse.get("name")) or "г. Барановичи"

    doc = Document()

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(BRAND_TITLE)
    header_run.bold = True
    header_run.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("АКТ О СПИСАНИИ ИНВЕНТАРЯ")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph(f"Дата составления: {writeoff_date}")
    doc.add_paragraph(f"Место составления: {place}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Наименование инструмента: {model_name}")
    doc.add_paragraph(f"Инвентарный номер: {inventory_number}")
    doc.add_paragraph(f"Серийный номер: {serial_number}")
    doc.add_paragraph(f"Причина списания: {reason}")
    doc.add_paragraph("")
    doc.add_paragraph("Подписи:")
    doc.add_paragraph("")
    doc.add_paragraph("Председатель комиссии: _________________________ /_________________/")
    doc.add_paragraph("")
    doc.add_paragraph("Член комиссии:         _________________________ /_________________/")
    doc.add_paragraph("")
    doc.add_paragraph("Материально ответственное лицо: ________________ /_________________/")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer


def _word_response(buffer: BytesIO, filename: str) -> StreamingResponse:
    buffer.seek(0)
    safe_filename = filename or "akt_spisaniya.docx"
    return StreamingResponse(
        buffer,
        media_type=DOCX_MEDIA,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_filename)}"},
    )


@router.get("/export/word/write-off/{tool_id}")
def export_writeoff_word(
    tool_id: UUID,
    supabase: Annotated[Client, Depends(get_supabase_client)],
    current_user: Annotated[CurrentUser, Depends(require_report_access)],
) -> StreamingResponse:
    """Акт списания инструмента в формате Word."""
    tool = _fetch_scrapped_tool(supabase, tool_id)

    if current_user.role == "clerk" and current_user.warehouse_id:
        if str(tool.get("warehouse_id")) != str(current_user.warehouse_id):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Доступ к инструменту другого склада запрещён",
            )

    buffer = _build_writeoff_docx(tool)
    inv_part = _cell(tool.get("inventory_number")).replace(" ", "_")[:40] or "tool"
    return _word_response(buffer, f"akt_spisaniya_{inv_part}.docx")


@router.get("/excel/inventory/{warehouse_id}")
def export_warehouse_inventory_legacy(
    warehouse_id: UUID,
    supabase: Annotated[Client, Depends(get_supabase_client)],
    current_user: Annotated[CurrentUser, Depends(require_report_access)],
) -> StreamingResponse:
    """Устаревший алиас экспорта остатков (совместимость)."""
    return export_inventory(warehouse_id, supabase, current_user)
